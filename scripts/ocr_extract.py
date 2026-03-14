"""
ocr_extract.py
Extracts plain text from a resume file (PDF or image).
- Text-based PDFs: uses pdfplumber → PyPDF2 fallback
- Scanned PDFs:    converts to images → PaddleOCR + Tesseract
- Images (JPG/PNG): PaddleOCR + Tesseract (complementary OCR for accuracy)

Outputs ONLY the extracted text to stdout.
All logs / warnings go to stderr (PHP only reads stdout).

Usage:
    python ocr_extract.py <file_path>
"""

import sys
import os
import logging

# ── Silence all logging before any heavy imports ────────────────────────────
logging.disable(logging.CRITICAL)
os.environ.setdefault('GLOG_logtostderr', '0')
os.environ.setdefault('GLOG_minloglevel', '3')
os.environ.setdefault('FLAGS_logtostderr', '0')
# Suppress paddle/paddleocr progress output
os.environ.setdefault('PADDLE_DISABLE_SIGNAL_HANDLER', '1')


# ────────────────────────────────────────────────────────────────────────────

def ocr_tesseract(file_path: str) -> str:
    """Run Tesseract OCR on an image file as complementary OCR engine."""
    try:
        import pytesseract
        text = pytesseract.image_to_string(file_path, lang='eng')
        return text.strip()
    except Exception:
        return ''


def ocr_image(file_path: str) -> str:
    """Run PaddleOCR + Tesseract on an image file for complementary accuracy; handles both 2.x and 3.x APIs."""
    paddle_text = ''
    tesseract_text = ''
    
    # Try PaddleOCR first (typically faster and good for structured text)
    try:
        from paddleocr import PaddleOCR
        ocr    = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        result = ocr.ocr(file_path, cls=True)
        if result:
            parts = []
            for page in result:
                if page is None:
                    continue
                if isinstance(page, list):
                    # API 2.x / compatible: [[bbox, [text, conf]], ...]
                    for line in page:
                        try:
                            word_info = line[1]
                            if isinstance(word_info, (list, tuple)):
                                parts.append(str(word_info[0]))
                            else:
                                parts.append(str(word_info))
                        except Exception:
                            pass
                elif hasattr(page, 'rec_texts'):
                    # PaddleOCR 3.x paddlex backend
                    parts.extend(page.rec_texts or [])
                elif hasattr(page, '__iter__'):
                    # Other iterable result objects
                    for item in page:
                        if hasattr(item, 'rec_texts'):
                            parts.extend(item.rec_texts or [])
            paddle_text = '\n'.join(str(p) for p in parts if p)
    except Exception:
        pass
    
    # Try Tesseract as complementary OCR (better for certain fonts/layouts)
    tesseract_text = ocr_tesseract(file_path)
    
    # Return PaddleOCR result if good; otherwise try Tesseract
    if paddle_text and len(paddle_text.strip()) > 50:
        return paddle_text
    elif tesseract_text and len(tesseract_text.strip()) > 50:
        return tesseract_text
    elif paddle_text:
        return paddle_text
    elif tesseract_text:
        return tesseract_text
    
    return ''


def pdf_text(file_path: str) -> str:
    """Extract text from a text-based PDF without OCR."""
    # Try pdfplumber (best quality)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = [p.extract_text() or '' for p in pdf.pages]
        text = '\n'.join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    # Try PyPDF2 / pypdf
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages  = [page.extract_text() or '' for page in reader.pages]
        text = '\n'.join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    return ''


def scanned_pdf_text(file_path: str) -> str:
    """Convert PDF pages to PNG images and run OCR on each."""
    try:
        from pdf2image import convert_from_path
        import tempfile

        pages = convert_from_path(file_path, dpi=200)
        parts = []
        for page in pages:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                page.save(tmp.name, 'PNG')
                tmp_path = tmp.name
            page_text = ocr_image(tmp_path)
            if page_text:
                parts.append(page_text)
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
        return '\n'.join(parts)
    except Exception:
        return ''


def docx_text(file_path: str) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx
        doc   = docx.Document(file_path)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return '\n'.join(parts)
    except Exception as e:
        print(str(e), file=sys.stderr)
        return ''


if __name__ == '__main__':
    if len(sys.argv) < 2:
        sys.exit(0)

    path = sys.argv[1]
    if not os.path.exists(path):
        sys.exit(0)

    ext  = os.path.splitext(path)[1].lower()
    text = ''

    if ext == '.pdf':
        text = pdf_text(path)
        if not text.strip():
            text = scanned_pdf_text(path)   # scanned PDF fallback
    elif ext == '.docx':
        text = docx_text(path)
    else:
        text = ocr_image(path)

    sys.stdout.write(text)
    sys.stdout.flush()
